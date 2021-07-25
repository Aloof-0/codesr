from docx import Document

# %cd D:YanZan_python2018word
document = Document()
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸


def word(w1, d1, d2, d3, w3, t1, t2, t3):
    # todo: 字体
    document.styles['Normal'].font.name = u'宋体'
    # todo:设置文档的基础样式
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # todo: 全局设置行边距
    document.styles['Normal'].paragraph_format.line_spacing=1.5
    pl = document.add_paragraph()
    # todo:对齐方式为居中，没有这句话默认左对齐
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = pl.add_run('\n关于向%s反馈谈话了结情况的函' % w1)
    run1.font.name = '方正小标宋简体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
    run1.font.size = Pt(22)
    # 前后距离5磅
    pl.space_before = Pt(5)
    pl.space_after = Pt(5)

    # 初始化建立第二个自然段
    p2 = document.add_paragraph()
    run2 = p2.add_run(
        "\n %s同志（被谈话人):\n\t %s年%s月%s日，%s（线索承办单位部门）请你就有关问题进行谈话了解情况。\n\t经研究，%s（线索承办单位部门）对你在谈话中所作的说明予以采信，谈话问题予以了结。" % (
        w1, d1, d2, d3, w3, w3))
    # 对客户的称呼

    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)
    # 是否加粗
    run2.font.bold = False

    p3 = document.add_paragraph()
    run2 = p3.add_run("\n \n 中共佛山市禅城区纪委监委办公室 \n %s年%s月%s日" % (t1, t2, t3))
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)
    # 是否加粗
    run2.font.bold = False
    # 对齐方式为右对齐，没有这句话默认左对齐
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save('../data/15-谈话情况了结函.docx')


word("王贝贝", "2021", "5", "12", "中共审查调查室", "2020", "8", "3")
